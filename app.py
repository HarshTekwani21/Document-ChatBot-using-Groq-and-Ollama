import os
import streamlit as st
import sqlite3
from dotenv import load_dotenv
from langchain.storage import LocalFileStore
from langchain.globals import set_llm_cache
from langchain_core.documents.base import Document
from langchain.embeddings import CacheBackedEmbeddings
from langchain_community.embeddings import OllamaEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.text_splitter import CharacterTextSplitter
from langchain.chains import ConversationalRetrievalChain, LLMChain
from langchain.memory import ConversationBufferMemory
from langchain_community.chat_message_histories import ChatMessageHistory
from langchain_community.chat_models import ChatOllama
from langchain_community.cache import SQLiteCache
from langchain.prompts import PromptTemplate
from docx import Document as DocxDocument
import openpyxl
import csv
import fitz  # PyMuPDF for PDF handling
import random

# Load environment variables
load_dotenv()

# Function to create a new SQLite connection
def create_db_connection():
    conn = sqlite3.connect('chat_history.db')
    c = conn.cursor()
    c.execute('''
        CREATE TABLE IF NOT EXISTS chat_history (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            conversation_id TEXT,
            user_message TEXT,
            bot_response TEXT
        )
    ''')
    conn.commit()
    return conn

class ChatAgent:
    def __init__(self, model_name="llama3-8b-8192"):
        self.groq_chat = ChatOllama(model="phi3:latest", temperature=0.2, max_tokens=128, verbose=True, api_key=os.getenv('GROQ_API_KEY'))
        set_llm_cache(SQLiteCache("llm_cache.db"))
        self.chain = None
        self.suggestion_chain = self.setup_suggestion_chain()

    def setup_suggestion_chain(self):
        prompt_template = """
        Based on the previous question and answer:
        Previous Question: {previous_question}
        Previous Answer: {previous_answer}
        Data Context: {data_context}

        Generate a list of follow-up questions that are relevant to the previous question and answer, and the data context.
        """
        prompt = PromptTemplate(template=prompt_template, input_variables=["previous_question", "previous_answer", "data_context"])
        
        # Initialize the suggestion chain with Groq's LLM
        llm = ChatOllama(api_key=os.getenv('GROQ_API_KEY'), model="phi3:latest", temperature=0.2, max_tokens=128, verbose=True)
        return LLMChain(llm=llm, prompt=prompt)

    def generate_suggestions(self, previous_question, previous_answer, data_context):
        return self.suggestion_chain.run({
            "previous_question": previous_question,
            "previous_answer": previous_answer,
            "data_context": data_context
        })

    def load_documents(self, folder_path):
        texts = []
        metadatas = []

        for root, _, files in os.walk(folder_path):
            for file in files:
                file_path = os.path.join(root, file)
                extension = os.path.splitext(file_path)[1].lower()

                if extension == ".pdf":
                    text_chunks, meta_chunks = self.load_pdf(file_path)
                elif extension == ".txt":
                    text_chunks, meta_chunks = self.load_txt(file_path)
                elif extension == ".docx":
                    text_chunks, meta_chunks = self.load_docx(file_path)
                elif extension == ".xlsx":
                    text_chunks, meta_chunks = self.load_xlsx(file_path)
                elif extension == ".csv":
                    text_chunks, meta_chunks = self.load_csv(file_path)
                else:
                    raise ValueError(f"Unsupported file format: {extension}")

                texts.extend(text_chunks)
                metadatas.extend(meta_chunks)

        return texts, metadatas

    def load_pdf(self, filepath):
        pdf_document = fitz.open(filepath)
        texts = []
        metadatas = []

        for page_num in range(len(pdf_document)):
            page = pdf_document.load_page(page_num)
            text = page.get_text()
            texts.append(text)
            metadatas.append({"Sources": f"pdf-{page_num}"})

        pdf_document.close()
        return texts, metadatas

    def load_txt(self, filepath):
        with open(filepath, 'r', encoding='utf-8') as file:
            text = file.read()
        return [text], [{"Sources": "txt-0"}]

    def load_docx(self, filepath):
        doc = DocxDocument(filepath)
        texts = []
        metadatas = []

        for i, paragraph in enumerate(doc.paragraphs):
            texts.append(paragraph.text)
            metadatas.append({"Sources": f"docx-{i}"})


        return texts, metadatas

    def load_xlsx(self, filepath):
        wb = openpyxl.load_workbook(filepath)
        sheet = wb.active
        texts = []
        metadatas = []

        for row in sheet.iter_rows(values_only=True):
            row_text = " ".join(str(cell) for cell in row if cell is not None)
            texts.append(row_text)
            metadatas.append({"Sources": f"xlsx-{sheet.title}"})

        return texts, metadatas

    def load_csv(self, filepath):
        texts = []
        metadatas = []

        with open(filepath, 'r', encoding='utf-8') as file:
            csv_reader = csv.reader(file)
            for i, row in enumerate(csv_reader):
                row_text = " ".join(row)
                texts.append(row_text)
                metadatas.append({"Sources": f"csv-{i}"})

        return texts, metadatas

    def textChunk_to_docObj(self, texts, metadatas):
        documents = [Document(page_content=text_chunk) for text_chunk in texts]

        underlying_embeddings = OllamaEmbeddings(model="nomic-embed-text")
        store = LocalFileStore("./cache/")
        cached_embedder = CacheBackedEmbeddings.from_bytes_store(
            underlying_embeddings, store, namespace=underlying_embeddings.model
        )

        db = FAISS.from_documents(documents, cached_embedder)

        message_history = ChatMessageHistory()
        memory = ConversationBufferMemory(
            memory_key="chat_history",
            output_key="answer",
            chat_memory=message_history,
            return_messages=True,
        )

        self.chain = ConversationalRetrievalChain.from_llm(
            llm=self.groq_chat,
            chain_type="stuff",
            retriever=db.as_retriever(),
            memory=memory,
            return_source_documents=True,
        )

    def chat(self, query, conversation_id, max_retries=5):
        try:
            docs = self.chain.invoke(query)
            answer = docs["answer"]

            # Create a new SQLite connection
            conn = create_db_connection()
            c = conn.cursor()

            # Store the chat in the database
            c.execute('INSERT INTO chat_history (conversation_id, user_message, bot_response) VALUES (?, ?, ?)',
                      (conversation_id, query, answer))
            conn.commit()
            conn.close()

            return answer
        except Exception as e:
            return f"Error: {str(e)}"

import streamlit as st

def main():
    # Load custom CSS
    st.markdown(
        """
        <style>
        .fixed-title {
            position: fixed;
            top: 18px;
            left: 350px;
            width: 100%;
            background-color: white;
            padding: 10px;
            z-index: 1000;
        }
        .content {
            margin-top: 60px; /* Adjust this value based on your title's height */
        }
        .chat-container {
            position: fixed;
            top: 35px;
            left: 350px;
            width: 100%;
            background-color: white;
            padding: 10px;
            z-index: 1000;
        }
        .input-container {
            position: fixed;
            bottom: 0;
            width: calc(80% - 40px); /* Adjust based on padding and border */
            display: flex;
            flex-direction: row;
            align-items: center;
            padding: 10px;
            background-color: #ffffff;
            border-top: 1px solid #ddd;
        }
        .suggestion-container {
            display: flex;
            overflow-x: auto;
            padding: 10px;
            background-color: #f9f9f9;
            border: 1px solid #ddd;
            margin-bottom: 10px;
            white-space: nowrap;
        }
        .suggestion-container button {
            flex: 0 0 auto;
            margin-right: 10px;
        }
        </style>
        """,
        unsafe_allow_html=True
    )

    # HTML for the fixed title
    st.markdown('<div class="fixed-title"><h1>Document ChatBot</h1></div>', unsafe_allow_html=True)

    # Initialize session state keys
    if "conversations" not in st.session_state:
        st.session_state.conversations = {}
        st.session_state.current_conversation = 1
    if "chat_history" not in st.session_state:
        st.session_state.chat_history = []
    if "query" not in st.session_state:
        st.session_state.query = ""
    if "selected_question" not in st.session_state:
        st.session_state.selected_question = ""
    if "suggestions" not in st.session_state:
        st.session_state.suggestions = []

    # Sidebar for conversation history
    st.sidebar.title("Conversation History")
    new_conversation = st.sidebar.button("New Conversation")
    if new_conversation:
        conversation_name = f"Conversation {st.session_state.current_conversation}"
        st.session_state.conversations[conversation_name] = st.session_state.chat_history.copy()
        st.session_state.chat_history = []
        st.session_state.current_conversation += 1

    selected_conversation = st.sidebar.selectbox("Select Conversation", ["None"] + list(st.session_state.conversations.keys()))
    if selected_conversation != "None":
        st.session_state.chat_history = st.session_state.conversations[selected_conversation]

    # Hardcoded folder path
    folder_path = "C:\\Users\\Harsh Tekwani\\Desktop\\ChatBot\\train"

    if folder_path:
        if "agent" not in st.session_state:
            st.session_state.agent = ChatAgent()

        agent = st.session_state.agent
        texts, metadatas = agent.load_documents(folder_path)
        agent.textChunk_to_docObj(texts, metadatas)    

    # Real-time chat container
    chat_container = st.container()
    with chat_container:
        st.markdown('<div class="chat-container"><h2 style="margin-top: 0;">Chat Window</h2></div>', unsafe_allow_html=True)
        for chat in st.session_state.chat_history:
            if chat['user']:
                st.markdown(f'<div style="text-align: right; background-color: #e6f2ff; color: black; padding: 10px; border-radius: 10px; margin: 10px;">{chat["user"]}</div>', unsafe_allow_html=True)
            if chat['bot']:
               st.markdown(f'<div style="text-align: left; background-color: #d9ead3; color: black; padding: 10px; border-radius: 10px; margin: 10px;">{chat["bot"]}</div>', unsafe_allow_html=True)

    # Suggestions container
    suggestions_container = st.container()
    with suggestions_container:
        st.markdown('<div class="suggestion-container">', unsafe_allow_html=True)
        st.subheader("Suggestions")
        for suggestion in st.session_state.suggestions:
            if st.button(suggestion):
                st.session_state.selected_question = suggestion
                st.experimental_rerun()
    

    # Input and submit container at the bottom
    input_container = st.container()
    with input_container:
        st.markdown('<div class="input-container">', unsafe_allow_html=True)
        input_col, submit_col = st.columns([5, 1])  # Adjust the ratio as needed

        with input_col:
            query = st.text_input("Your message", value=st.session_state.selected_question, label_visibility="collapsed")

        with submit_col:
            if st.button("Submit"):
                if "agent" in st.session_state:
                    agent = st.session_state.agent
                    conversation_id = f"Conversation {st.session_state.current_conversation}"
                    answer = agent.chat(query, conversation_id)
                    st.session_state.chat_history.append({"user": query, "bot": answer})

                    # Generate and display suggestion questions
                    previous_question = query
                    previous_answer = answer
                    data_context = "Contextual information for suggestion generation."
                    suggestions = agent.generate_suggestions(previous_question, previous_answer, data_context)
                    st.session_state.suggestions = [suggestion.strip() for suggestion in suggestions.split('\n') if suggestion.strip()][:5]  # Limit to 5 suggestions
                    st.session_state.selected_question = ""  # Clear selected question

                    st.experimental_rerun()
                else:
                    st.warning("Please specify a folder path first.")
        st.markdown('</div>', unsafe_allow_html=True)

if __name__ == "__main__":
    main()
